"""DOCX export for portfolio rollup data.

Generates a landscape Word document with capability > epic > story tables.
"""
import io
import logging
from datetime import datetime, timezone
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


class ExportService:
    def generate(
        self,
        capabilities: List[Dict[str, Any]],
        view: str = "progress",
        filter_name: str = "all",
        project_key: str = "",
    ) -> bytes:
        """Generate DOCX from capability data.

        Returns bytes of the generated document.
        """
        doc = Document()

        # Set landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # Title
        doc.add_heading('Portfolio Rollup', level=1)
        now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        doc.add_paragraph(
            f"Project: {project_key} | View: {view.title()} | "
            f"Filter: {filter_name.title()} | Generated: {now}",
            style='Subtitle'
        )

        # Summary
        total_caps = len(capabilities)
        total_cum = sum(
            c.get("rollups", {}).get("cumulative_points", 0) for c in capabilities
        )
        total_rem = sum(
            c.get("rollups", {}).get("remaining_points", 0) for c in capabilities
        )
        doc.add_paragraph(
            f"{total_caps} capabilities \u00b7 {total_cum:.0f} cumulative pts "
            f"\u00b7 {total_rem:.0f} remaining pts"
        )

        # Per-capability sections
        for cap in capabilities:
            r = cap.get("rollups", {})
            doc.add_heading(
                f"{cap.get('key', '')} \u2014 {cap.get('summary', '')}",
                level=2
            )
            doc.add_paragraph(
                f"Status: {cap.get('status', '')} | "
                f"Cumulative: {r.get('cumulative_points', 0):.0f} | "
                f"Remaining: {r.get('remaining_points', 0):.0f}"
            )

            epics = cap.get("epics", [])
            if not epics:
                doc.add_paragraph("No epics.", style='List Bullet')
                continue

            # Table for epics
            if view == "progress":
                headers = ["Key", "Summary", "Status", "Size", "Cumul.", "Remain."]
            elif view == "schedule":
                headers = ["Key", "Summary", "Status", "Start", "End", "Days"]
            else:  # cycle
                headers = [
                    "Key", "Summary", "Status", "Dev", "QA", "Stage", "Prod",
                ]

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Header row
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)

            # Epic rows
            for epic in epics:
                er = epic.get("rollups", {})
                row = table.add_row()
                row.cells[0].text = epic.get("key", "")
                row.cells[1].text = epic.get("summary", "")[:60]
                row.cells[2].text = epic.get("status", "")

                if view == "progress":
                    size = epic.get("tshirt_size", "") or ""
                    if epic.get("uses_tshirt_fallback"):
                        size += " *"
                    row.cells[3].text = size
                    row.cells[4].text = str(
                        int(er.get("cumulative_points", 0))
                    )
                    row.cells[5].text = str(
                        int(er.get("remaining_points", 0))
                    )

                # Set font size for all cells
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

            doc.add_paragraph()  # spacing

        # Footer
        doc.add_paragraph(
            f"Generated {now} \u00b7 {total_caps} capabilities",
            style='Intense Quote',
        )

        # Serialize to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
